#!/usr/bin/env python3
"""
Build a single .docx with N PNGs that fit the A4 page width.
You specify the target total size (GB/GiB) and the script computes the PNG size.

- Uses python-docx (3rd party) for the .docx body
- Generates valid PNG bytes (pure stdlib) with predictable size (zlib level 0)
- Displays ALL images (one per paragraph), scaled to page width
- All PNGs share the same byte size; content is random (different seeds) but size is identical

Examples:
  python make_target_sized_docx.py --docx-path out.docx --num-images 2500 --target-size 3.9 --unit GB
  python make_target_sized_docx.py --docx-path out2.docx --num-images 1000 --target-size 2 --unit GB

Notes:
- For very large N (e.g., 10k) Word 2016 may open slowly. If you only need size (not to view all),
  I can give a variant that embeds all images but only shows the first few.
"""

import argparse
import io
import os
import random
import struct
import zlib
from typing import Tuple

from docx import Document  # third-party
from docx.shared import Cm  # for A4 + margins + width fit

# ---------- PNG generator (valid 8-bit RGB, predictable size) ----------
PNG_SIG = b"\x89PNG\r\n\x1a\n"


def _crc32(data: bytes) -> int:
    import binascii
    return binascii.crc32(data) & 0xFFFFFFFF


def _chunk(typ: bytes, data: bytes) -> bytes:
    return (struct.pack(">I", len(data)) + typ + data +
            struct.pack(">I", _crc32(typ + data)))


def build_png_bytes(width: int, height: int, seed: int) -> bytes:
    """
    Minimal 8-bit RGB PNG, filter type 0 per row, zlib level 0.
    Size ~= 8 (sig) + IHDR + IDAT(zlib hdr + raw) + IEND.
    Raw = height * (1 + 3*width), constant regardless of pixel randomness.
    """
    rng = random.Random(seed)
    # IHDR
    ihdr = struct.pack(">IIBBBBB",
                       width, height,
                       8,  # bit depth
                       2,  # color type (truecolor RGB)
                       0, 0, 0)
    # Raw data: for each row: 1 filter byte (0) + 3*width bytes of RGB
    row_len = 1 + 3 * width
    raw = bytearray(row_len * height)
    p = 0
    for _ in range(height):
        raw[p] = 0  # filter 0
        p += 1
        for _ in range(width * 3):
            raw[p] = rng.randrange(0, 512)
            p += 1
    # zlib compress with level 0 (stored blocks): predictable size ~ raw + small overhead
    idat = zlib.compress(bytes(raw), level=0)
    png = bytearray()
    png += PNG_SIG
    png += _chunk(b'IHDR', ihdr)
    png += _chunk(b'IDAT', idat)
    png += _chunk(b'IEND', b'')
    return bytes(png)


def choose_png_geometry_for_size(per_image_target: int, width: int = 512) -> Tuple[int, bytes]:
    """
    Given a desired per-image size (bytes), compute a height that yields >= target.
    We keep width fixed (512) and increase height until we reach the target.
    Returns (height, png_bytes).
    """
    base_overhead = 100  # cushion
    row_bytes = 1 + 3 * width
    height = max(1, (per_image_target - base_overhead) // row_bytes)
    seed = 1234
    png = build_png_bytes(width, height, seed)
    attempts = 0
    while len(png) < per_image_target and attempts < 20000:
        height += 1
        seed += 1
        png = build_png_bytes(width, height, seed)
        attempts += 1
    return height, png


# ---------- DOCX builder ----------
def make_docx(docx_path: str, num_images: int, target_total_bytes: int) -> dict:
    """
    Create one .docx with num_images PNGs displayed, scaled to A4 text width.
    Compute per-image target size from total target; generate one geometry and reuse it
    (using different seeds so pixels differ while the size stays identical).
    """
    if num_images < 1:
        raise ValueError("num_images must be >= 1")

    # Reserve a small overhead for XML/ZIP central directory (~ a few hundred KB).
    overhead_cushion = 600_000
    target_for_images = max(1, target_total_bytes - overhead_cushion)
    per_image_target = max(1, target_for_images // num_images)

    # Find a PNG geometry (constant) that meets per-image target
    height, sample_png = choose_png_geometry_for_size(per_image_target, width=512)
    per_png_bytes = len(sample_png)

    # === Build the Word doc ===
    doc = Document()
    # A4 page + reasonable margins (2 cm)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Content width in cm = page_width - margins
    content_width_cm = (section.page_width - section.left_margin - section.right_margin) / Cm(1)

    doc.add_paragraph(
        f"Generated DOCX with {num_images} PNGs (each ~{per_png_bytes:,} bytes). "
        f"All images are scaled to page width ({content_width_cm:.2f} cm)."
    )

    # Insert all images visible, scaled to content width
    # Reuse geometry; vary seed to make different pixels while keeping size identical
    seed = 10000
    for i in range(1, num_images + 1):
        png_bytes = build_png_bytes(512, height, seed)
        seed += 1
        # sanity: enforce identical size (for determinism)
        if len(png_bytes) != per_png_bytes:
            # extremely unlikely with fixed geometry and level=0; but if it happens, adjust
            png_bytes = sample_png
        stream = io.BytesIO(png_bytes)
        doc.add_picture(stream, width=Cm(content_width_cm))
        # optional caption/paragraph separator
        # doc.add_paragraph(f"Image {i}")

    doc.save(docx_path)

    final_size = os.path.getsize(docx_path)
    return {
        "docx_path": docx_path,
        "png_width": 512,
        "png_height": height,
        "per_png_bytes": per_png_bytes,
        "target_bytes": target_total_bytes,
        "final_bytes": final_size,
    }


# ---------- CLI ----------
def to_bytes(value: float, unit: str) -> int:
    if unit.lower() == "gib":
        return int(value * (1024 ** 3))
    return int(value * 1_000_000_000)  # GB decimal


def from_bytes(n: int, unit: str) -> float:
    if unit.lower() == "gib":
        return n / (1024 ** 3)
    return n / 1_000_000_000


def main():
    ap = argparse.ArgumentParser(description="Make one .docx of ~target size with N PNGs fitting A4 width.")
    ap.add_argument("--docx-path", required=True, help="Output .docx path")
    ap.add_argument("--num-images", type=int, required=True, help="How many PNGs to embed & display")
    ap.add_argument("--target-size", type=float, required=True, help="Target size value")
    ap.add_argument("--unit", choices=["GB", "GiB"], default="GB", help="Unit for target size")
    args = ap.parse_args()

    target_bytes = to_bytes(args.target_size, args.unit)
    info = make_docx(args.docx_path, args.num_images, target_bytes)

    fb = info["final_bytes"]
    print(f"Created: {info['docx_path']}")
    print(f"Images: {args.num_images} | PNG geometry: {info['png_width']}x{info['png_height']} px")
    print(f"Per PNG bytes: {info['per_png_bytes']:,}")
    print(f"Target: ~{args.target_size} {args.unit}  |  Actual: {fb:,} bytes "
          f"(≈ {from_bytes(fb, 'GB'):.6f} GB | {from_bytes(fb, 'GiB'):.6f} GiB)")
    print("Note: If you need closer to the exact target, bump --target-size slightly; "
          "ZIP/XML overhead can vary by a few hundred KB.")


if __name__ == "__main__":
    main()